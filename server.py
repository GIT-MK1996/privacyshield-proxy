import os
import json
import re
import zipfile
import difflib
import anthropic
from http.server import HTTPServer, BaseHTTPRequestHandler
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn

PORT = int(os.environ.get("PORT", 8080))
API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
client = anthropic.Anthropic(api_key=API_KEY)

PROMPT = """Je bent een AVG-privacytool. Vervang ALLE persoonsgegevens in de tekst door labels.

Vervang:
- Persoonsnamen (ook initialen zoals W.J.G. Riel, T. Hellfayer) → <PERSOON>
- Bedrijfsnamen (ook bv, B.V., nv, VOF etc.) → <BEDRIJF>
- Woon- en geboorteplaatsen → <PLAATS>
- Personeels/klantnummers → <ID-NUMMER>
- Leeftijd als herleidbaar → <LEEFTIJD>

Laat deze labels zoals ze zijn:
<IBAN>, <BSN>, <E-MAIL>, <TELEFOON>, <ADRES>, <POSTCODE>, <BEDRAG>, <DATUM>, <TIJD>, <IP-ADRES>, <KENTEKEN>, <ID-NUMMER>, <GEBRUIKERS-ID>, <WERKSTATION>, <BESTANDSNAAM>

REGELS:
- Vervang ELKE voorkomen, ook als naam 10x voorkomt
- Geef UITSLUITEND de geanonimiseerde tekst terug, geen uitleg
- Verander niks aan de rest van de tekst

Tekst:
"""

PATTERNS = [
    (re.compile(r'\bNL\d{2}[A-Z]{4}\d{10}\b'), '<IBAN>'),
    (re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'), '<E-MAIL>'),
    (re.compile(r'(?:\+31|0031|0)[- ]?(?:6[- ]?\d{8}|[1-9]\d[- ]?\d{7}|\d{2}[- ]?\d{6,7})'), '<TELEFOON>'),
    (re.compile(r'(?<![A-Z0-9\-])\b\d{9}\b(?![A-Z0-9\-])'), '<BSN>'),
    (re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b'), '<IP-ADRES>'),
    (re.compile(r'\b\d{4}\s?[A-Z]{2}\b'), '<POSTCODE>'),
    (re.compile(r'\b[A-Z][a-z]+(?:straat|laan|weg|plein|gracht|dijk|singel|kade|dreef|steeg|pad|allee|boulevard|hof|markt|dam|brug|wende)\s+\d+[a-zA-Z]?\b'), '<ADRES>'),
    (re.compile(r'(?:€|EUR)\s*\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?'), '<BEDRAG>'),
    (re.compile(r'\d{1,2}[\-/\.]\d{1,2}[\-/\.]\d{2,4}'), '<DATUM>'),
    (re.compile(r'\b\d{1,2}\s+(?:januari|februari|maart|april|mei|juni|juli|augustus|september|oktober|november|december)\s+\d{4}\b', re.IGNORECASE), '<DATUM>'),
    (re.compile(r'\b\d{1,2}:\d{2}(?:\s?uur)?\b'), '<TIJD>'),
    (re.compile(r'\b[A-Z]{1,3}-\d{1,3}-[A-Z0-9]{1,3}\b'), '<KENTEKEN>'),
    (re.compile(r'\b[A-Z]{1,6}-\d{4}-\d{3,6}\b'), '<ID-NUMMER>'),
    (re.compile(r'\buser_[a-z0-9_]{3,20}\b', re.IGNORECASE), '<GEBRUIKERS-ID>'),
]

def vaste_patronen(tekst):
    for patroon, label in PATTERNS:
        tekst = patroon.sub(label, tekst)
    return tekst

def anonymize_with_claude(tekst):
    tussen = vaste_patronen(tekst)
    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        messages=[{"role": "user", "content": PROMPT + tussen}]
    )
    return message.content[0].text

def has_drawing(para):
    return bool(para._p.findall('.//' + qn('w:drawing')))

def anonymize_docx(file_bytes):
    """Verwerk DOCX: bewaar logo en afbeeldingen, vervang alleen tekst."""
    doc = Document(BytesIO(file_bytes))

    # Verzamel alle tekst voor Claude
    alle_tekst = []
    for para in doc.paragraphs:
        if not has_drawing(para) and para.text.strip():
            alle_tekst.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not has_drawing(para) and para.text.strip():
                        alle_tekst.append(para.text)

    # Claude anonimiseert alles in één keer
    gecombineerd = '\n'.join(alle_tekst)
    geanonimiseerd = anonymize_with_claude(gecombineerd)

    # Bouw vervangingen-woordenboek via difflib
    vervangingen = {}
    for orig, anon in zip(gecombineerd.split('\n'), geanonimiseerd.split('\n')):
        if orig != anon:
            matcher = difflib.SequenceMatcher(None, orig, anon)
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == 'replace' and i2 - i1 > 1:
                    vervangingen[orig[i1:i2]] = anon[j1:j2]

    print(f"Vervangingen gevonden: {len(vervangingen)}")

    # Pas toe per run — runs nooit samenvoegen
    def process_para(para):
        if has_drawing(para):
            return
        if not para.text.strip():
            return
        for run in para.runs:
            if not run.text:
                continue
            nieuw = run.text
            for patroon, label in PATTERNS:
                nieuw = patroon.sub(label, nieuw)
            for orig, anon in vervangingen.items():
                nieuw = nieuw.replace(orig, anon)
            if nieuw != run.text:
                run.text = nieuw

    for para in doc.paragraphs:
        process_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)

    # Sla op met python-docx
    docx_output = BytesIO()
    doc.save(docx_output)
    docx_output.seek(0)
    docx_bytes = docx_output.read()

    # Haal originele media op (logo, handtekeningen etc.)
    orig_media = {}
    with zipfile.ZipFile(BytesIO(file_bytes), 'r') as zorig:
        for name in zorig.namelist():
            if 'media' in name:
                orig_media[name] = zorig.read(name)

    # Zet originele media terug in het nieuwe bestand
    output = BytesIO()
    with zipfile.ZipFile(BytesIO(docx_bytes), 'r') as znew:
        with zipfile.ZipFile(output, 'w') as zout:
            for item in znew.infolist():
                if item.filename in orig_media:
                    zout.writestr(item, orig_media[item.filename], compress_type=item.compress_type)
                else:
                    zout.writestr(item, znew.read(item.filename), compress_type=item.compress_type)

    output.seek(0)
    return output.read()


class Handler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        print(f"{self.command} {self.path} - {args[1]}")

    def send_cors(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_cors()
        self.end_headers()

    def do_GET(self):
        if self.path == "/health":
            self.send_response(200)
            self.send_cors()
            self.send_header("Content-Type", "application/json")
            self.end_headers()
            self.wfile.write(json.dumps({"status": "ok"}).encode())
        else:
            self.send_response(404)
            self.end_headers()

    def do_POST(self):
        length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(length)

        if self.path == "/anonimiseer":
            try:
                data = json.loads(body)
                tekst = data.get("tekst", "")
                if not tekst:
                    raise ValueError("Geen tekst")
                result = anonymize_with_claude(tekst)
                self.send_response(200)
                self.send_cors()
                self.send_header("Content-Type", "application/json")
                self.end_headers()
                self.wfile.write(json.dumps({"result": result}).encode())
            except Exception as e:
                print(f"Fout: {e}")
                self.send_response(500)
                self.send_cors()
                self.send_header("Content-Type", "application/json")
                self.end_headers()
                self.wfile.write(json.dumps({"error": str(e)}).encode())

        elif self.path == "/anonimiseer-docx":
            try:
                result_bytes = anonymize_docx(body)
                self.send_response(200)
                self.send_cors()
                self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                self.send_header("Content-Disposition", "attachment; filename=geanonimiseerd.docx")
                self.send_header("Content-Length", str(len(result_bytes)))
                self.end_headers()
                self.wfile.write(result_bytes)
            except Exception as e:
                print(f"DOCX fout: {e}")
                self.send_response(500)
                self.send_cors()
                self.send_header("Content-Type", "application/json")
                self.end_headers()
                self.wfile.write(json.dumps({"error": str(e)}).encode())

        else:
            self.send_response(404)
            self.end_headers()


if __name__ == "__main__":
    print(f"PrivacyShield Python proxy op poort {PORT}")
    print(f"API key aanwezig: {bool(API_KEY)}")
    server = HTTPServer(("0.0.0.0", PORT), Handler)
    server.serve_forever()
